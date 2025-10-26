# extract_actas.py
# --------------------------------------------------------
# Funciones para procesar PDFs con Google Document AI
# (usando credenciales desde Streamlit Secrets)
# + fallback local para PDF/DOCX cuando haga falta.
# --------------------------------------------------------

import io
import streamlit as st
import pandas as pd

from google.cloud import documentai_v1 as documentai
from google.oauth2 import service_account

from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document as DocxDocument


# ============= CREDENCIALES =============
def get_gcp_credentials(scopes=None):
    """
    Carga credenciales desde Streamlit Secrets.
    Debes haber pegado tu JSON en:
      [gcp_service_account]
      ...
    """
    creds = service_account.Credentials.from_service_account_info(
        st.secrets["gcp_service_account"]
    )
    if scopes:
        creds = creds.with_scopes(scopes)
    return creds


# ============= DOCUMENT AI =============
def process_with_document_ai(file_bytes: bytes, mime_type: str = "application/pdf"):
    """
    Procesa un archivo con tu Custom Extractor de Document AI.
    Lee project/location/processor desde [docai] en secrets.
    Devuelve: (texto_completo, dataframe_de_entidades)
    """
    project_id  = st.secrets["docai"]["project_id"]
    location    = st.secrets["docai"]["location"]
    processor_id= st.secrets["docai"]["processor_id"]

    creds = get_gcp_credentials()
    client = documentai.DocumentProcessorServiceClient(credentials=creds)
    name = f"projects/{project_id}/locations/{location}/processors/{processor_id}"

    raw_document = documentai.RawDocument(content=file_bytes, mime_type=mime_type)
    request = documentai.ProcessRequest(name=name, raw_document=raw_document)
    result = client.process_document(request=request)
    doc = result.document

    # Texto completo del documento
    full_text = doc.text or ""

    # Entidades etiquetadas por tu Custom Extractor
    rows = []
    for e in doc.entities or []:
        page = None
        try:
            if e.page_anchor and e.page_anchor.page_refs:
                page = e.page_anchor.page_refs[0].page
        except Exception:
            page = None

        rows.append({
            "Etiqueta": e.type_,
            "Valor": e.mention_text,
            "Confianza": round(getattr(e, "confidence", 0.0), 3),
            "Página": page,
        })

    df_entities = pd.DataFrame(rows)
    return full_text, df_entities


# ============= FALLBACK LOCAL =============
def extract_text_local(uploaded_file):
    """
    Si Document AI no está disponible, extrae texto local:
    - PDF con pdfminer.six
    - DOCX con python-docx
    - TXT como texto plano
    """
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        # pdfminer necesita un path o un buffer binario
        data = uploaded_file.read()
        buffer = io.BytesIO(data)
        text = pdf_extract_text(buffer)
        return text

    if name.endswith(".docx"):
        data = uploaded_file.read()
        buffer = io.BytesIO(data)
        doc = DocxDocument(buffer)
        text = "\n".join(p.text for p in doc.paragraphs)
        return text

    # Texto plano u otros
    return uploaded_file.read().decode("utf-8", errors="ignore")
